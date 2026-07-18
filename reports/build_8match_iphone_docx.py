from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("Football_AI_Pro_8_matches_full_analysis_2026-07-17.docx")

NAVY = "1F4D78"
INK = "202A33"
MUTED = "64748B"
GREEN = "2F6B4F"
GREEN_FILL = "EAF4EE"
AMBER = "8A6100"
AMBER_FILL = "FFF6DD"
RED = "9A3030"
RED_FILL = "FCEBEC"


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(paragraph, fill, border=None):
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    ppr.append(node)
    if border:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        borders.append(left)
        ppr.append(borders)


def line(doc, label, value, *, fill=None, accent=None, value_color=INK, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if fill:
        p.paragraph_format.left_indent = Inches(0.08)
        p.paragraph_format.right_indent = Inches(0.08)
        shade(p, fill, accent)
    a = p.add_run(label)
    set_font(a, 10.8, True, accent or NAVY)
    b = p.add_run(value)
    set_font(b, 10.8, False, value_color)
    return p


def heading(doc, value, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(value)
    set_font(r, {1: 16, 2: 13, 3: 12}[level], True, NAVY)
    return p


def note(doc, title, text, fill, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.25
    shade(p, fill, color)
    r = p.add_run(title + "\n")
    set_font(r, 11.5, True, color)
    r = p.add_run(text)
    set_font(r, 10.8)


def add_match(doc, item):
    heading(doc, f"{item['no']}  {item['name']}", 2)
    line(doc, "处理：", item["status"], fill=item["fill"], accent=item["color"])
    line(doc, "竞彩胜平负：", item["one_x_two"])
    line(doc, "让球胜平负：", item["handicap"])
    line(doc, "盘口信息：", item["line"])
    line(doc, "AI / 信心 / 冷门：", f"{item['ai']} / {item['confidence']} / {item['upset']}")
    line(doc, "比分参考：", item["score"])
    line(doc, "判断依据：", item["reason"], after=9)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

# compact_reference_guide design tokens
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, before, after in [
    ("Heading 1", 16, 18, 10),
    ("Heading 2", 13, 14, 7),
    ("Heading 3", 12, 10, 5),
]:
    st = doc.styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(NAVY)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("Football AI Pro | 赛前分析汇总"), 8.5, False, MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("数据快照：2026-07-17 | 仅供研究与赛前复核"), 8, False, MUTED)

# memo_masthead, no decorative rule
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("Football AI Pro"), 12, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("8 场比赛完整分析汇总"), 24, True, NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(14)
set_font(p.add_run("竞彩胜平负 + 让球胜平负 + AI评分 + 赛果路径 + 串关方案"), 11.5, False, MUTED)

note(
    doc,
    "阅读说明",
    "本报告整合用户提供的竞彩、球探多公司欧赔与亚盘截图。AI评分表示数据一致性；信心指数表示方向相对把握，不是保证命中率或实际胜率。",
    AMBER_FILL,
    AMBER,
)

heading(doc, "一、最终筛选与放弃标注", 1)
note(doc, "正式通过", "204 巴伊亚主胜；207 纳什维尔主胜；202 米亚尔比主胜。", GREEN_FILL, GREEN)
note(doc, "条件通过", "201 哥德堡主胜（防平）；205 弗鲁米嫩塞主胜（不作稳胆）。", AMBER_FILL, AMBER)
note(doc, "建议放弃", "203 博德闪耀（仅开 -2 深盘）；206 米拉索尔（市场强度分歧）；208 洛城德比（波动过高）。", RED_FILL, RED)

heading(doc, "二、8 场比赛逐场完整结论", 1)
matches = [
    dict(no="201", name="瑞超：哥德堡 vs 布鲁马波卡纳", status="条件通过", fill=AMBER_FILL, color=AMBER,
         one_x_two="主胜，防平", handicap="-1：让负为主，防让平", line="竞彩：主胜 1.84 / 平 3.50 / 主负 3.33；让球 -1", ai=72, confidence=68, upset=35,
         score="1-0 / 1-1 / 2-1", reason="主队是第一方向，但多家公司亚洲盘主要停留在主让半球，未形成强势升盘；可单关观察，不宜充当串关稳胆。"),
    dict(no="202", name="瑞超：米亚尔比 vs 韦斯特罗斯", status="正式通过（稳胆候选）", fill=GREEN_FILL, color=GREEN,
         one_x_two="主胜", handicap="-1：让平优先，防让负", line="竞彩：主胜 1.61 / 平 3.75 / 主负 4.15；让球 -1", ai=74, confidence=70, upset=28,
         score="1-0 / 2-1", reason="主场优势明确，标准盘主胜可做稳胆。模型比分集中于一球小胜，因此让球 -1 的“让平”比穿盘更贴合高赔路径。"),
    dict(no="203", name="挪超：博德闪耀 vs 腓特烈斯塔", status="建议放弃", fill=RED_FILL, color=RED,
         one_x_two="标准盘未开售，不纳入组合", handicap="-2：让平或让负", line="竞彩仅见让球 -2：主胜 1.90 / 平 4.10 / 主负 2.78", ai=80, confidence=76, upset=18,
         score="2-0 / 3-1", reason="主胜方向虽强，但可投市场被迫进入 -2 深盘；赢球不等同赢盘，净胜两球是结算边界，风险不适合作串关。"),
    dict(no="204", name="巴甲：巴伊亚 vs 沙佩科恩斯", status="正式通过（第一稳胆）", fill=GREEN_FILL, color=GREEN,
         one_x_two="主胜", handicap="-1：让胜，防让平", line="竞彩：主胜 1.29 / 平 4.70 / 主负 7.10；让球 -1", ai=83, confidence=81, upset=12,
         score="2-0 / 2-1", reason="多家公司欧赔与亚盘对主队形成一致支撑，且竞彩主胜低位稳定；是本组方向最清晰的一场，标准主胜优先。"),
    dict(no="205", name="巴甲：弗鲁米嫩塞 vs 布拉干RB", status="条件通过", fill=AMBER_FILL, color=AMBER,
         one_x_two="主胜，防平", handicap="-1：让平或让负", line="竞彩：主胜 1.69 / 平 3.34 / 主负 4.20；让球 -1", ai=68, confidence=62, upset=31,
         score="1-0 / 2-1 / 1-1", reason="主胜为第一方向，但竞彩主胜强度高于市场平均强度，且让球盘并不支持大胜；可作备选，不能归入稳胆。"),
    dict(no="206", name="巴甲：米拉索尔 vs 格雷米奥", status="建议放弃", fill=RED_FILL, color=RED,
         one_x_two="主胜，防平", handicap="-1：让负", line="竞彩：主胜 1.72 / 平 3.20 / 主负 4.25；让球 -1", ai=65, confidence=59, upset=39,
         score="1-0 / 1-1", reason="竞彩对主队的定价明显强于多家公司市场均值，三项概率接近，市场一致性不足；不满足正式组合门槛。"),
    dict(no="207", name="美职：纳什维尔 vs 亚特联", status="正式通过（第二稳胆）", fill=GREEN_FILL, color=GREEN,
         one_x_two="主胜", handicap="-1：让胜或让平", line="竞彩：主胜 1.34 / 平 4.45 / 主负 6.30；让球 -1", ai=81, confidence=78, upset=15,
         score="2-0 / 2-1", reason="主胜低位与多家公司欧赔同步，客胜处于普遍高位。主胜标准盘是更稳的选择；让球盘可赢可走边界，单选风险更高。"),
    dict(no="208", name="美职：洛城银河 vs 洛杉矶FC", status="建议放弃", fill=RED_FILL, color=RED,
         one_x_two="客胜，防平", handicap="+1：让胜或让平", line="竞彩：主胜 2.95 / 平 3.55 / 主负 1.97；让球 +1", ai=67, confidence=60, upset=48,
         score="1-2 / 1-1 / 2-2", reason="客胜是市场第一方向，但同城德比的波动与平局权重均高，盘口仅至客让平半，无法支持强客胜结论；建议放弃。"),
]
for item in matches:
    add_match(doc, item)

heading(doc, "三、修正后的 3 串 1（2 个稳胆标准盘 + 1 个高赔让球盘）", 1)
note(doc, "主推组合（按截图赔率约 5.88 倍）", "巴伊亚 主胜 @1.29 × 纳什维尔 主胜 @1.34 × 米亚尔比（-1）让球平 @3.40。", GREEN_FILL, GREEN)
line(doc, "组合逻辑：", "前两场只用标准盘主胜做方向稳胆；第三场不用强行追主胜，而选米亚尔比一球小胜最匹配的 -1 让球平，既保留模型路径，也把组合赔率提升到目标区间。")
note(doc, "备选组合（按截图赔率约 6.40 倍）", "巴伊亚 主胜 @1.29 × 纳什维尔 主胜 @1.34 × 哥德堡（-1）让球平 @3.70。", AMBER_FILL, AMBER)
line(doc, "备选风险：", "哥德堡属于条件通过，盘面没有提供同等强度的支持；所以该组合仅作更高波动备选，优先级低于主推组合。")
note(doc, "组合门禁", "203、206、208 不进入正式串关；205 仅作替补，不替代前两场稳胆。", RED_FILL, RED)

heading(doc, "四、风险提示与临场复核", 1)
for text in [
    "赔率会继续变化。只有出现“盘口升降 + 多家公司同步变化 + 基本面支持”时，才提高临场信号权重。",
    "竞彩胜平负与让球胜平负是不同结算市场。强队赢球，并不自动意味着让球盘能赢。",
    "如临场出现核心伤停、轮换、首发与预期明显不一致，需重新评估串关门槛。",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), 10.8)

doc.core_properties.title = "Football AI Pro - 8场比赛完整分析汇总"
doc.core_properties.subject = "竞彩胜平负、让球胜平负与串关方案"
doc.core_properties.author = "Football AI Pro"
doc.save(OUT)
print(OUT)
