from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r'C:\Users\apple\Documents\足彩分析\reports\2026-07-17_十场重新筛选_盘口列修正版.docx')

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(9)
    if color: r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc = Document(); sec = doc.sections[0]; sec.top_margin = Inches(.65); sec.bottom_margin = Inches(.65); sec.left_margin = Inches(.7); sec.right_margin = Inches(.7)
styles = doc.styles; styles['Normal'].font.name = 'Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei'); styles['Normal'].font.size = Pt(10)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = title.add_run('Football AI Pro｜十场比赛重新筛选报告'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(31,78,121)
sub = doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.add_run('概率校准 + 弃权门禁 + 多公司一致性复核\n数据截点：用户提供截图（2026-07-17 15:19–15:28）').italic=True

p=doc.add_paragraph(); rr=p.add_run('重要说明：'); rr.bold=True; p.add_run('本报告依据已提供截图和已登记联赛样本重新筛选。AI评分、信心指数仅表示当前数据一致性，不是保证命中率。标记“弃”代表不进入正式组合。')

doc.add_heading('一、最终筛选总表', level=1)
doc.add_paragraph('盘口标注说明： “竞彩让球”列为官方让球数值；-1 表示主队让一球，+1 表示主队受让一球。该列与“让球胜平负”配套阅读。')
headers=['场次','对阵','竞彩让球','竞彩胜平负','让球胜平负','AI评分','信心','处理']
rows=[
('201','瓦勒伦加 vs 奥勒松','-1','主胜','让平',72,68,'条件通过'),('202','德里城 vs 索陆军','+1','客胜','让平',73,67,'条件通过'),('203','费伦茨瓦罗斯 vs 伏伊伏丁','-1','主胜','让胜',82,77,'通过·最强'),('204','日利纳 vs 斯海杜克','+1','客胜','让平',68,62,'观望'),('205','博塔弗戈 vs 桑托斯','-1','平','让负',63,57,'弃'),('206','维多利亚 vs 达伽马','-1','主胜防平','让平',64,57,'弃'),('207','蒙特利尔 vs 多伦多','-1','平','让负',62,56,'弃'),('208','芝加哥 vs 温哥华','+1','平','让胜',64,59,'弃'),('209','圣路易斯城 vs 堪萨斯城','-1','主胜','让平',75,69,'通过'),('210','西雅图 vs 波特兰','-1','主胜','让平',74,68,'通过')]
t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
for i,h in enumerate(headers): set_cell(t.rows[0].cells[i],h,True,(255,255,255)); shade(t.rows[0].cells[i],'1F4E79')
for row in rows:
    cells=t.add_row().cells
    for i,v in enumerate(row):
        color=(192,0,0) if v=='弃' else None; set_cell(cells[i],v, i==6, color)
        if v=='通过·最强': shade(cells[i],'E2F0D9')
        elif v=='弃': shade(cells[i],'FCE4D6')

doc.add_heading('二、保留场次', level=1)
for title_txt, body in [
('203｜费伦茨瓦罗斯 vs 伏伊伏丁','标准盘主胜为十场中最强方向。多家公司欧赔同步压低主胜，亚洲盘出现跨档升盘，联赛样本与盘口方向一致。竞彩建议：主胜；让球建议：让胜，但让胜风险高于普通主胜。'),
('209｜圣路易斯城 vs 堪萨斯城','标准盘主胜方向稳定，主场样本支持主队不败。亚洲盘虽强化主队，但受让方仍有低水保护，不能把升盘直接等同于大胜。竞彩建议：主胜；让球建议：让平。'),
('210｜西雅图 vs 波特兰','主胜概率明确，欧赔和基本面均支持主队。让球盘对受让方存在保护，净胜两球的不确定性较高。竞彩建议：主胜；让球建议：让平。'),
('201｜瓦勒伦加 vs 奥勒松','主胜方向成立，但防守样本不够稳定，属于条件通过。竞彩建议：主胜；让球建议：让平，不进入最稳层。'),
('202｜德里城 vs 索陆军','客胜是欧赔第一方向，但平局保护明显，客胜优势不够大。竞彩建议：客胜；让球建议：让平，临场客胜赔率明显上升时退出。'),
('204｜日利纳 vs 斯海杜克','客胜方向有支持，但跨公司一致性和净胜优势不足。建议仅作观察，不进入核心串关。')]:
    p=doc.add_paragraph(); p.style='Heading 2'; p.add_run(title_txt)
    doc.add_paragraph(body)

doc.add_heading('三、弃权场次', level=1)
for txt in [
'205｜博塔弗戈 vs 桑托斯：欧赔与亚洲盘方向冲突，平局概率高，模型与市场缺乏一致性。',
'206｜维多利亚 vs 达伽马：主胜赔率与盘口水位不匹配，主胜边际不足；联赛样本只支持主队轻微优势。',
'207｜蒙特利尔 vs 多伦多：主胜、平局和客队受让方向互相牵制，无法形成可执行优势。',
'208｜芝加哥 vs 温哥华：欧赔偏向客队，但盘口对主队有保护，属于典型冲突盘。'
]: doc.add_paragraph(txt, style='List Bullet')

doc.add_heading('四、组合建议', level=1)
doc.add_paragraph('稳健方案：203主胜 × 209主胜。只选两场，优先命中率。')
doc.add_paragraph('高风险3串1（参考赔率约9.46倍）：203主胜（-1）让胜 @2.11 × 209主胜（-1）普通主胜 @1.32 × 202主胜（+1）让平 @3.40。该组合不作为首选，风险主要来自203让胜与202让平。')
doc.add_paragraph('稳健3串1：203普通主胜 @1.35 × 209普通主胜 @1.32 × 202客胜 @1.66，参考组合赔率约2.96倍，命中率优先但回报较低。')
doc.add_paragraph('折中3串1：203普通主胜 @1.35 × 209普通主胜 @1.32 × 202（+1）让平 @3.40，参考组合赔率约6.06倍；这是本报告更符合“两个稳胆+一个中高风险腿”的组合。')
doc.add_paragraph('不建议把205、206、207、208加入任何正式串关。')

doc.add_heading('五、临场退出条件', level=1)
for txt in ['203若临场退回浅盘且主队水位持续升高，取消让胜，仅保留普通主胜。','202若客胜赔率明显升至1.85以上，退出稳健层。','209、210若让球继续升档但受让方持续低水，维持让平，不升级让胜。','任何单一公司变化，不足以改变结论；至少三家公司同步且方向一致才提高权重。']:
    doc.add_paragraph(txt, style='List Bullet')

doc.add_heading('六、门禁结论', level=1)
doc.add_paragraph('本次十场筛选结果：4场正式候选、2场条件候选、4场弃权。弃权不是预测其必败，而是表示当前数据不足以支持正式下注。该策略会降低覆盖率，但符合提高命中率的目标。')
doc.save(OUT)
print(OUT)
