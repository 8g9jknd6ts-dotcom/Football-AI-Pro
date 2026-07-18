from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r'C:\Users\apple\Documents\足彩分析')
REPORTS = ROOT / 'reports'
REPORTS.mkdir(exist_ok=True)

NAVY = '183B56'; BLUE = '2F6690'; PALE = 'EAF2F8'; LIGHT = 'F6F8FA'; GREEN = 'E7F3EC'; AMBER = 'FFF4D6'; RED = 'FBE7E7'; INK = '203040'; MUTED = '657786'

matches = [
    ('201','瓦勒伦加 vs 奥勒松','-1','主胜','让平',72,68,'条件通过'),
    ('202','德里城 vs 索陆军','+1','客胜','让平',73,67,'条件通过'),
    ('203','费伦茨瓦罗斯 vs 伏伊伏丁','-1','主胜','让胜',82,77,'通过·最强'),
    ('204','日利纳 vs 斯海杜克','+1','客胜','让平',68,62,'观望'),
    ('205','博塔弗戈 vs 桑托斯','-1','平','让负',63,57,'弃'),
    ('206','维多利亚 vs 达伽马','-1','主胜防平','让平',64,57,'弃'),
    ('207','蒙特利尔 vs 多伦多','-1','平','让负',62,56,'弃'),
    ('208','芝加哥 vs 温哥华','+1','平','让胜',64,59,'弃'),
    ('209','圣路易斯城 vs 堪萨斯城','-1','主胜','让平',75,69,'通过'),
    ('210','西雅图 vs 波特兰','-1','主胜','让平',74,68,'通过'),
]

reasons = {
 '203':'标准盘主胜为十场中最强方向；多家公司同步压低主胜，亚盘出现升盘，联赛样本与盘口方向一致。让胜风险高于普通主胜。',
 '209':'标准盘主胜方向较稳；主场样本支持不败。亚盘虽偏主队，但受让方仍有低水保护，因此让平优先。',
 '210':'欧赔和基本面均支持主队；亚盘对受让方有保护，净胜两球不确定，竞彩让球以让平为主。',
 '201':'主胜方向成立，但防守稳定性不足，只列条件通过。',
 '202':'客胜为欧赔首选，但平局保护明显，临场客胜升水则退出稳健层。',
 '204':'客胜有支持但公司间一致性不足，先观望。',
 '205':'欧赔与盘口方向冲突，平局风险高，弃。',
 '206':'主胜赔率与盘口水位不完全匹配，弃。',
 '207':'主胜、平局、受让方向互相牵制，弃。',
 '208':'欧赔偏客队但让球保护主队，冲突明显，弃。',
}

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell(cell, text, size=9, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    r = p.add_run(str(text)); r.font.name = 'Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei'); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        g = OxmlElement('w:gridCol'); g.set(qn('w:w'), str(w)); grid.append(g)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.first_child_found_in('w:tcW')
            if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa'); set_cell_margins(cell)

def base_doc(phone=False):
    d=Document(); sec=d.sections[0]
    if phone:
        sec.page_width=Inches(6.0); sec.page_height=Inches(9.0); sec.top_margin=Inches(.45); sec.bottom_margin=Inches(.45); sec.left_margin=Inches(.42); sec.right_margin=Inches(.42)
    else:
        sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    st=d.styles['Normal']; st.font.name='Microsoft YaHei'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(10 if phone else 10.5); st.font.color.rgb=RGBColor.from_string(INK); st.paragraph_format.space_after=Pt(5); st.paragraph_format.line_spacing=1.12
    for name,size,col,before,after in [('Heading 1',15,NAVY,12,6),('Heading 2',12,BLUE,8,4),('Heading 3',11,NAVY,6,3)]:
        s=d.styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    return d

def title_block(d, phone):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(2); r=p.add_run('Football AI Pro｜十场比赛筛选报告'); r.bold=True; r.font.name='Microsoft YaHei'; r.font.size=Pt(17 if phone else 20); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(7); r=p.add_run('概率校准 · 弃权门禁 · 多公司一致性复核\n数据截点：2026-07-17 15:19–15:28'); r.italic=True; r.font.size=Pt(8.5 if phone else 9.5); r.font.color.rgb=RGBColor.from_string(MUTED)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(7); r=p.add_run('阅读提示：'); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE); p.add_run('AI评分和信心指数表示当前数据一致性，不等同于保证命中率。竞彩让球列中，-1=主队让一球，+1=主队受让一球。')

def add_phone(d):
    title_block(d, True)
    d.add_heading('一、结论先看',1)
    p=d.add_paragraph(); p.paragraph_format.space_after=Pt(8); r=p.add_run('首选核心：203 主胜；209 主胜；210 主胜。\n折中3串1：203主胜 × 209主胜 × 202（+1）让平，参考赔率约6.06倍。'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    d.add_heading('二、十场快速筛选',1)
    table=d.add_table(rows=1, cols=5); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'; set_widths(table,[620,2640,700,850,950])
    for c,h in zip(table.rows[0].cells,['场次','对阵','让球','竞彩方向','处理']): set_cell(c,h,8.5,True,'FFFFFF'); shade(c,NAVY)
    for no,fixture,hcap,spf,ah,score,conf,status in matches:
        cells=table.add_row().cells; vals=[no,fixture,hcap,spf,status]
        fill=GREEN if '通过' in status else (AMBER if status=='观望' else (RED if status=='弃' else LIGHT))
        for c,v in zip(cells,vals): set_cell(c,v,8.3, v==status, INK, WD_ALIGN_PARAGRAPH.CENTER); shade(c,fill if v==status else 'FFFFFF')
    d.add_heading('三、保留场次卡片',1)
    for no in ['203','209','210','201','202','204']:
        row=next(x for x in matches if x[0]==no); no,fixture,hcap,spf,ah,score,conf,status=row
        p=d.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2); r=p.add_run(f'{no}｜{fixture}'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
        p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.space_after=Pt(2); p.add_run(f'竞彩让球 {hcap}  ·  胜平负 {spf}  ·  让球胜平负 {ah}  ·  AI {score}  ·  信心 {conf}  ·  {status}')
        p=d.add_paragraph(reasons[no]); p.paragraph_format.left_indent=Inches(.08); p.paragraph_format.space_after=Pt(5); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
    d.add_heading('四、弃权场次',1)
    d.add_paragraph('205、206、207、208：盘口与欧赔/基本面出现冲突，当前一致性不足，不进入正式串关。')
    d.add_heading('五、组合建议',1)
    d.add_paragraph('稳健2串1：203主胜 × 209主胜，参考赔率约1.78倍。')
    d.add_paragraph('折中3串1（首选）：203主胜 × 209主胜 × 202（+1）让平，参考赔率约6.06倍。')
    d.add_paragraph('高风险3串1：203（-1）让胜 × 209主胜 × 202（+1）让平，参考赔率约9.46倍，不作为首选。')
    d.add_heading('六、临场门禁',1); d.add_paragraph('仅当至少三家公司同向变化、盘口升降与欧赔方向一致、且无重大首发/伤停反转时，才保留原组合。')

def add_desktop(d):
    title_block(d, False)
    d.add_heading('一、最终筛选总表',1)
    table=d.add_table(rows=1, cols=8); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'; set_widths(table,[560,2100,650,850,850,650,650,1050])
    for c,h in zip(table.rows[0].cells,['场次','对阵','竞彩让球','胜平负','让球胜平负','AI评分','信心','处理']): set_cell(c,h,8.5,True,'FFFFFF'); shade(c,NAVY)
    for no,fixture,hcap,spf,ah,score,conf,status in matches:
        cells=table.add_row().cells; vals=[no,fixture,hcap,spf,ah,score,conf,status]; fill=GREEN if '通过' in status else (AMBER if status=='观望' else (RED if status=='弃' else LIGHT))
        for i,(c,v) in enumerate(zip(cells,vals)): set_cell(c,v,8.5, i==7, INK); shade(c,fill if i==7 else 'FFFFFF')
    d.add_heading('二、重点场次与依据',1)
    for no in ['203','209','210','201','202','204']:
        row=next(x for x in matches if x[0]==no); _,fixture,hcap,spf,ah,score,conf,status=row
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(2); r=p.add_run(f'{no}｜{fixture}'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
        d.add_paragraph(f'竞彩让球：{hcap}；胜平负：{spf}；让球胜平负：{ah}；AI评分：{score}；信心：{conf}；处理：{status}。{reasons[no]}')
    d.add_heading('三、弃权门禁',1)
    d.add_paragraph('205 博塔弗戈 vs 桑托斯、206 维多利亚 vs 达伽马、207 蒙特利尔 vs 多伦多、208 芝加哥 vs 温哥华：盘口与欧赔/基本面冲突，标记“弃”，不进入正式串关。')
    d.add_heading('四、组合建议',1)
    d.add_paragraph('稳健2串1：203主胜 × 209主胜，参考赔率约1.78倍。')
    d.add_paragraph('折中3串1（首选）：203主胜 × 209主胜 × 202（+1）让平，参考赔率约6.06倍；符合“两稳胆+一中高风险腿”。')
    d.add_paragraph('高风险3串1：203（-1）让胜 × 209主胜 × 202（+1）让平，参考赔率约9.46倍，不作为首选。')
    d.add_heading('五、临场门禁与风险',1)
    d.add_paragraph('单一公司水位变化不提高权重；至少三家公司同向、盘口升降与欧赔方向一致，并且无重大首发/伤停反转，才维持原结论。AI评分/信心是数据一致性指标，不是命中率承诺。')

phone=base_doc(True); add_phone(phone); phone.save(REPORTS/'2026-07-17_十场筛选_iPhone优化版.docx')
desktop=base_doc(False); add_desktop(desktop); desktop.save(REPORTS/'2026-07-17_十场筛选_电脑明细优化版.docx')
print('created', REPORTS/'2026-07-17_十场筛选_iPhone优化版.docx')
print('created', REPORTS/'2026-07-17_十场筛选_电脑明细优化版.docx')
