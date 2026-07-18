from pathlib import Path
import sqlite3
from datetime import datetime

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "Football_AI_Pro_8Match_Postmortem_20260719.docx"
DB = ROOT / "data" / "football_ai_tracking.sqlite"

MATCHES = [
    dict(no="201", mid="20260718-SWE-IFKGOT-BP", league="瑞超", game="哥德堡 vs 布鲁马波卡纳", score="2-1", std="主胜", jline="主让-1", jresult="让平", asia="主队 -0.5：赢", prior="标准主胜、防平；让球首选让负，防让平", verdict="标准判断命中；让球防守方向命中", ai="72 / 68", note="主让半没有升盘，赛前就不应把穿盘当作主结论。2-1正好落在让平，是‘赢球不穿’的典型。"),
    dict(no="202", mid="20260718-SWE-MJALLBY-VASTERAS", league="瑞超", game="米亚尔比 vs 韦斯特罗", score="0-0", std="平", jline="主让-1", jresult="让负", asia="主队 -0.5：输", prior="标准主胜；让球首选让平，防让负", verdict="核心标准判断失误；让球备选覆盖", ai="74 / 70", note="主胜低赔未兑现，0-0说明进攻兑现风险被低估；该场不应进入稳胆层。"),
    dict(no="203", mid="20260718-NOR-BODO-FREDRIKSTAD", league="挪超", game="博德闪耀 vs 腓特烈斯塔", score="1-0", std="主胜", jline="主让-2", jresult="让负", asia="主队 -2/-2.5：输", prior="不做标准盘；深让-2取让平/让负路径，弃赛", verdict="弃赛门禁正确；让球路径覆盖", ai="80 / 76", note="只靠点球1-0。赛后信息显示世界杯后轮换/疲劳存在，深盘与阵容负荷不匹配，弃赛判断比猜胜负更有价值。"),
    dict(no="204", mid="20260718-BRA-BAHIA-CHAPECOENSE", league="巴甲", game="巴伊亚 vs 沙佩科恩斯", score="2-0", std="主胜", jline="主让-1", jresult="让胜", asia="主队 -0.75/-1：赢", prior="标准主胜；让球胜，防让平", verdict="标准、让球、亚洲方向均命中", ai="83 / 81", note="本轮最完整的一场：强势主队、市场共识和盘口兑现一致；可作为稳胆样本，但仍不能外推为保证。"),
    dict(no="205", mid="20260718-BRA-FLUMINENSE-BRAGANTINO", league="巴甲", game="弗鲁米嫩塞 vs 布拉干RB", score="1-1", std="平", jline="主让-1", jresult="让负", asia="主队 -0.5：输", prior="标准主胜、防平；让球平/负，条件通过", verdict="主胜失误；防平与让球路径正确", ai="68 / 62", note="赛前已给主胜防平，实际被平局击中。低一致性场不应按‘可串关’处理。"),
    dict(no="206", mid="20260718-BRA-MIRASSOL-GREMIO", league="巴甲", game="米拉索尔 vs 格雷米奥", score="2-1", std="主胜", jline="主让-1", jresult="让平", asia="主队 -0.5：赢", prior="标准主胜、防平；让负；弃赛", verdict="方向正确，但弃赛仍合理", ai="65 / 59", note="主胜兑现但只赢一球。由于赛前市场分歧和低信心，不能因赛后结果反向证明当时应重仓。"),
    dict(no="207", mid="20260718-MLS-NASHVILLE-ATLANTA", league="美职联", game="纳什维尔 vs 亚特联", score="1-0", std="主胜", jline="主让-1", jresult="让平", asia="主队 -1：走水", prior="标准主胜；让球胜/平", verdict="标准命中；让球路径命中（亚洲走水）", ai="81 / 78", note="主胜兑现但只赢一球，说明标准盘稳胆不等于让一球必穿；亚洲-1按走水处理，不能记为赢。"),
    dict(no="208", mid="20260718-MLS-LAG-LAFC", league="美职联", game="洛城银河 vs 洛杉矶FC", score="0-3", std="客胜", jline="主受+1", jresult="让负", asia="主队 +0.25：输", prior="标准客胜、防平；让球胜/平；弃赛", verdict="标准方向命中；让球路径失误", ai="67 / 60", note="德比波动被正确标注为弃赛，但最终客队大胜。赛后不能把正确结果等同为可复制的赛前可下注信号。"),
]

SOURCES = [
    ("201", "IFK Göteborg 官方赛后报道 / 瑞典媒体赛况", "2-1；客队33分钟红牌，主队末段绝杀信息已核验"),
    ("202", "Aftonbladet / 442 赛后报道", "0-0；主队占优未能破门"),
    ("203", "VG 赛后报道", "1-0；点球唯一进球，世界杯后轮换与疲劳背景"),
    ("204", "Esporte Clube Bahia 官方 / CBF", "巴伊亚2-0沙佩科恩斯"),
    ("205", "Fluminense 官方 / CBF", "弗鲁米嫩塞1-1布拉干RB，末段追平"),
    ("206", "Grêmio 官方赛后报告", "米拉索尔2-1格雷米奥"),
    ("207", "Nashville SC 官方赛后报告", "纳什维尔1-0亚特联，联赛不败延续"),
    ("208", "LA Galaxy / LAFC 官方赛后报告", "银河0-3洛杉矶FC；射正2-6、控球59.5%-40.5%"),
]

def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def set_cell(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text)); r.bold = bold; r.font.name = 'Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei'); r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    for i,h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, True, 'FFFFFF', 9); shade(table.rows[0].cells[i], '1F4E78')
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row): set_cell(cells[i], val, size=8.5)
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def title(doc, text, level=1):
    p=doc.add_paragraph(); p.style=f'Heading {level}'; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(text); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.color.rgb=RGBColor(31,78,120)

def body(doc, text, bold_prefix=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.15
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); r.bold=True; p.add_run(text[len(bold_prefix):])
    else: p.add_run(text)
    for r in p.runs:
        r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); r.font.size=Pt(9.5)

def setup(doc):
    sec=doc.sections[0]; sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5); sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    styles=doc.styles
    styles['Normal'].font.name='Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); styles['Normal'].font.size=Pt(9.5)
    for s in ['Heading 1','Heading 2']:
        styles[s].font.name='Microsoft YaHei'; styles[s]._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rr=header.add_run('Football AI Pro  |  赛后复盘'); rr.font.size=Pt(8); rr.font.color.rgb=RGBColor(100,100,100)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    f=footer.add_run('数据依据：用户提供的赛前球探/竞彩截图 + 赛后官方/权威报道；仅作数据复盘，不构成投注承诺。'); f.font.size=Pt(7.5); f.font.color.rgb=RGBColor(110,110,110)

def update_db():
    schema=(ROOT/'data'/'schema.sql').read_text(encoding='utf-8')
    with sqlite3.connect(DB) as conn:
        conn.executescript(schema)
        now='2026-07-19T10:00:00+08:00'
        for m in MATCHES:
            h,a=map(int,m['score'].split('-'))
            result='H' if h>a else 'A' if h<a else 'D'
            conn.execute('''INSERT OR REPLACE INTO matches(match_id,kickoff_utc,league_code,league_name,home_team,away_team,source_name,source_match_ref,source_snapshot_path,source_reliability,created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?)''',(m['mid'],None,m['league'],m['league'],m['game'].split(' vs ')[0],m['game'].split(' vs ')[1],'User screenshots + verified reports',m['no'],'reports/Football_AI_Pro_8Match_Postmortem_20260719.docx','high',now))
            conn.execute('''INSERT OR REPLACE INTO outcomes(match_id,recorded_at,actual_home_goals,actual_away_goals,actual_1x2,actual_handicap,actual_total_goals,result_source,result_source_ref,verified)
            VALUES(?,?,?,?,?,?,?,?,?,1)''',(m['mid'],now,h,a,result,m['jresult'],h+a,'Official club/league report + reputable match report',m['no']))
        conn.execute('''INSERT OR REPLACE INTO user_preferences(preference_id,recorded_at,category,preference_text,source,status,applied_to_version,notes)
        VALUES(?,?,?,?,?,?,?,?)''',('Pref002',now,'reporting','赛后复盘必须逐场呈现赛前建议、赛果、竞彩标准盘/让球盘结算、亚洲盘结算、错误归因及串关盈亏；手机端以单场卡片和窄表呈现。','user','active','Football AI Pro 3.1','2026-07-19 user request'))
        conn.commit()

def build():
    doc=Document(); setup(doc)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Football AI Pro'); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(31,78,120); r.font.name='Georgia'
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('8场比赛赛后全面复盘'); r.bold=True; r.font.size=Pt(18); r.font.name='Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei')
    p=doc.add_paragraph('覆盖：2026-07-18 赛前报告（场次201—208）  |  复盘日：2026-07-19'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size=Pt(9); p.runs[0].font.color.rgb=RGBColor(90,90,90)
    title(doc,'一、结论先看')
    body(doc,'赛前正式三项“主方案”中，巴伊亚主胜、纳什维尔主胜命中；米亚尔比主胜失误，主方案3串1未中。替代方案中的哥德堡-1让平命中，因此替代3串1按约6.40倍结算为赢。')
    body(doc,'核心复盘：本轮最有价值的并非“猜对了多少主胜”，而是深盘弃赛规则有效挡住了博德闪耀-2的落空；但米亚尔比被误列为稳胆，说明低赔主胜仍需经过进球兑现门禁。')
    add_table(doc,['层级','赛前选择','实际','结论'],[
        ['主方案3串1','巴伊亚主胜 × 纳什维尔主胜 × 米亚尔比(-1)让平','前两项命中；米亚尔比让负','未命中，1单位=-1'],
        ['替代3串1','巴伊亚主胜 × 纳什维尔主胜 × 哥德堡(-1)让平','三项均命中','约6.40倍，1单位净+5.40'],
        ['弃赛门禁','博德闪耀深让-2、米拉索尔、洛城德比','博德-2让负；其余有方向但波动大','弃赛原则正确，不以赛后结果倒推重仓'],
    ],[3.0,5.2,5.2,3.0])
    title(doc,'二、总览：赛前与赛后结算')
    rows=[]
    for m in MATCHES:
        rows.append([m['no'],m['game'],m['score'],m['std'],m['jline']+' '+m['jresult'],m['verdict']])
    add_table(doc,['场次','比赛','赛果','标准盘','竞彩让球','复盘结论'],rows,[1.0,3.9,1.1,1.4,2.4,5.2])
    body(doc,'统计口径：仅把赛前明确写出的标准盘方向纳入方向验证。7场有明确标准盘首选的比赛中，主/客/平单选方向命中5场；但这不是模型正式命中率，更不能与收益率混同。让球盘中存在“首选+防守”双路径，故只列结算，不把双选覆盖伪装成单选命中率。')
    title(doc,'三、逐场复盘（手机阅读版）')
    for m in MATCHES:
        title(doc,f"{m['no']}｜{m['league']}｜{m['game']}",2)
        add_table(doc,['赛前判断','实际结算','结论'],[[m['prior'],f"赛果 {m['score']}；标准盘：{m['std']}；竞彩{m['jline']}：{m['jresult']}；亚洲：{m['asia']}",m['verdict']]], [5.6,6.1,4.2])
        body(doc,'AI评分 / 信心指数：'+m['ai']+'。说明：这两个值只反映赛前数据的一致性与相对把握，不是承诺胜率。')
        body(doc,'复盘要点：'+m['note'])
    title(doc,'四、盘口与模型诊断')
    body(doc,'1）“标准盘主胜”与“让一球穿盘”必须拆开：哥德堡、纳什维尔、米拉索尔均主胜但竞彩-1仅让平；不能因主胜判断正确就给让球盘加分。')
    body(doc,'2）低赔稳胆需新增进球兑现门禁：米亚尔比0-0说明市场虽给出主队优势，却未证明其至少两球/至少一球的兑现能力。后续对主胜低于1.70且让球盘为-1的场次，若缺少明确火力、首发或连续多公司同步支撑，不进入稳胆串关。')
    body(doc,'3）深盘与阵容负荷联动：博德闪耀虽1-0赢球，但-2失守。世界杯后复赛、轮换、疲劳和高让球并存时，深盘自动降级或弃赛。')
    body(doc,'4）弃赛不能被赛后结果否定：米拉索尔主胜、LAFC客胜都实现了方向，但赛前属于分歧/德比高波动场。模型应该奖励“避免无法证明优势的入场”，不是赛后追认。')
    title(doc,'五、规则库更新（待累计样本回测后才可升为正式模型）')
    add_table(doc,['编号','候选规则','本轮依据','状态'],[
        ['Rule041','低赔主胜进球兑现门禁','米亚尔比0-0，低赔主胜未兑现','候选；待联赛样本回测'],
        ['Rule042','深盘×轮换/国际赛事负荷降权','博德闪耀1-0，-2让负','候选；待跨联赛回测'],
        ['Rule043','标准盘与让球盘独立记分','多场主胜但-1让平/让负','已执行为报告口径'],
        ['Rule044','弃赛结果不得反向计入推荐命中','米拉索尔、洛城德比赛后方向正确','已执行为复盘口径'],
    ],[1.7,5.0,6.2,3.0])
    title(doc,'六、赛果核验来源')
    for no,src,note in SOURCES: body(doc,f'{no}｜{src}：{note}')
    body(doc,'限制说明：赛前盘口以用户提供的竞彩/球探截图为准；亚洲盘按截图所示主流盘口结算。未保存逐分钟完整赔率轨迹的比赛，不把赛后水位变化解释为确定因果。')
    OUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUT)

if __name__=='__main__':
    update_db(); build(); print(OUT)
